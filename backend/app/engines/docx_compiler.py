"""
Word Compiler — converts DocumentAST into a .docx file using python-docx.
"""
from pathlib import Path
from typing import Optional
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from app.engines.document_ast import (
    DocumentAST,
    Section,
    Paragraph,
    TextRun,
    CitationRun,
    FigureRef,
    TableRef,
    EquationRef,
    FigureElement,
    TableElement,
    AuthorInfo,
)


class DocxCompiler:
    """
    Compiles a DocumentAST into a .docx file.
    Supports formatting rules from journal templates.
    """

    def __init__(self, template_dir: str = "../templates"):
        self.template_dir = Path(template_dir)

    def compile(
        self,
        ast: DocumentAST,
        output_path: str,
        formatting_rules: Optional[dict] = None,
    ) -> str:
        """
        Compile DocumentAST to .docx file.

        Args:
            ast: The DocumentAST to compile.
            output_path: Where to save the .docx file.
            formatting_rules: Optional dict with font, spacing, margin specs.

        Returns:
            The output file path.
        """
        rules = formatting_rules or {}

        doc = Document()

        # ---- Page setup ----
        for section in doc.sections:
            section.top_margin = Cm(float(rules.get("margin_top_cm", 2.54)))
            section.bottom_margin = Cm(float(rules.get("margin_bottom_cm", 2.54)))
            section.left_margin = Cm(float(rules.get("margin_left_cm", 2.54)))
            section.right_margin = Cm(float(rules.get("margin_right_cm", 2.54)))

        # ---- Default style ----
        style = doc.styles["Normal"]
        font_name = rules.get("font_family", "Times New Roman")
        font_size = int(rules.get("font_size", 12))
        style.font.name = font_name
        style.font.size = Pt(font_size)
        line_spacing = float(rules.get("line_spacing", 2.0))
        style.paragraph_format.line_spacing = line_spacing

        # ---- Title ----
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(ast.title)
        title_run.bold = True
        title_run.font.size = Pt(font_size + 4)

        # ---- Authors ----
        if ast.authors:
            author_para = doc.add_paragraph()
            author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            author_names = ", ".join(
                f"{a.first_name} {a.last_name}" for a in ast.authors
            )
            author_para.add_run(author_names)

            if ast.authors[0].affiliation:
                aff_para = doc.add_paragraph()
                aff_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                aff_run = aff_para.add_run(ast.authors[0].affiliation)
                aff_run.font.size = Pt(font_size - 1)
                aff_run.italic = True

        doc.add_paragraph()  # spacer

        # ---- Abstract ----
        if ast.abstract:
            heading = doc.add_heading("Abstract", level=1)
            self._apply_heading_style(heading, font_name, font_size)
            for para in ast.abstract.paragraphs:
                p = doc.add_paragraph()
                self._compile_paragraph_to_docx(p, para, ast)

        # ---- Sections ----
        for section in ast.sections:
            heading = doc.add_heading(section.heading, level=1)
            self._apply_heading_style(heading, font_name, font_size)
            for para in section.paragraphs:
                p = doc.add_paragraph()
                p.style.font.name = font_name
                p.style.font.size = Pt(font_size)
                p.paragraph_format.line_spacing = line_spacing
                self._compile_paragraph_to_docx(p, para, ast)

        # ---- Figures ----
        if ast.figures:
            doc.add_heading("Figures", level=1)
            for fig in ast.figures:
                # Try to insert image if file exists
                try:
                    if Path(fig.image_path).exists():
                        doc.add_picture(fig.image_path, width=Inches(5.5))
                except Exception:
                    pass
                cap_para = doc.add_paragraph()
                cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap_run = cap_para.add_run(
                    f"Fig. {fig.number}. {fig.caption}"
                )
                cap_run.italic = True

        # ---- Tables ----
        if ast.tables:
            doc.add_heading("Tables", level=1)
            for tbl in ast.tables:
                cap_para = doc.add_paragraph()
                cap_para.add_run(
                    f"Table {tbl.number}. {tbl.caption}"
                ).bold = True

                num_rows = len(tbl.rows) + (1 if tbl.headers else 0)
                num_cols = len(tbl.headers) if tbl.headers else (
                    len(tbl.rows[0]) if tbl.rows else 1
                )
                table = doc.add_table(rows=num_rows, cols=num_cols)
                table.style = "Table Grid"

                if tbl.headers:
                    for j, h in enumerate(tbl.headers):
                        cell = table.rows[0].cells[j]
                        cell.text = h
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
                    start_row = 1
                else:
                    start_row = 0

                for i, row_data in enumerate(tbl.rows):
                    for j, cell_text in enumerate(row_data):
                        if j < num_cols:
                            table.rows[start_row + i].cells[j].text = str(cell_text)

                doc.add_paragraph()  # spacer

        # ---- References ----
        if ast.references:
            doc.add_heading("References", level=1)
            for ref in ast.references:
                ref_para = doc.add_paragraph()
                ref_para.add_run(f"[{ref.number}] ").bold = True
                ref_para.add_run(ref.formatted_text)

        # ---- Save ----
        doc.save(output_path)
        return output_path

    def _compile_paragraph_to_docx(self, p, para: Paragraph, ast: DocumentAST):
        """Add runs from a Paragraph to a python-docx paragraph object."""
        for run in para.runs:
            if isinstance(run, TextRun):
                r = p.add_run(run.text)
                r.bold = run.bold
                r.italic = run.italic
                if run.subscript:
                    r.font.subscript = True
                if run.superscript:
                    r.font.superscript = True
            elif isinstance(run, CitationRun):
                # Resolve citation numbers
                ref_nums = self._resolve_citation_numbers(run.ref_ids, ast)
                p.add_run(f" [{', '.join(map(str, ref_nums))}]")
            elif isinstance(run, FigureRef):
                # Find figure number
                fig_num = self._resolve_figure_number(run.figure_label, ast)
                p.add_run(f" Fig. {fig_num}")
            elif isinstance(run, TableRef):
                tbl_num = self._resolve_table_number(run.table_label, ast)
                p.add_run(f" Table {tbl_num}")
            elif isinstance(run, EquationRef):
                eq_num = self._resolve_equation_number(run.equation_label, ast)
                p.add_run(f" Eq. {eq_num}")

    @staticmethod
    def _apply_heading_style(heading, font_name: str, base_size: int):
        """Apply consistent heading formatting."""
        for run in heading.runs:
            run.font.name = font_name
            run.font.size = Pt(base_size + 2)

    @staticmethod
    def _resolve_citation_numbers(ref_ids, ast: DocumentAST):
        """Map cite_keys to citation numbers."""
        key_to_num = {r.cite_key: r.number for r in ast.references}
        return [key_to_num.get(rid, 0) for rid in ref_ids]

    @staticmethod
    def _resolve_figure_number(label: str, ast: DocumentAST) -> int:
        """Map a figure label to its number."""
        for fig in ast.figures:
            if fig.label == label:
                return fig.number
        return 0

    @staticmethod
    def _resolve_table_number(label: str, ast: DocumentAST) -> int:
        for tbl in ast.tables:
            if tbl.label == label:
                return tbl.number
        return 0

    @staticmethod
    def _resolve_equation_number(label: str, ast: DocumentAST) -> int:
        for eq in ast.equations:
            if eq.label == label:
                return eq.number
        return 0
